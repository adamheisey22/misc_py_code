import os
from docx import Document
import pandas as pd

def extract_feedback_from_folder(folder_path):
    """
    Extracts all responses from multiple .docx files in a folder, after the questions,
    into a single DataFrame, and tags each response with the file name's 'name' part.
    Assumes file names follow the format: 'form_g1_name.docx'
    """
    all_responses = []
    
    # Iterate over all files in the folder
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            # Extract the 'name' part from the file name (e.g., 'form_g1_heisey.docx' -> 'heisey')
            name_tag = file_name.split('_')[-1].replace('.docx', '')
            
            # Load the document
            doc_path = os.path.join(folder_path, file_name)
            doc = Document(doc_path)
            
            current_question = None
            responses = []
            
            # Iterate over paragraphs to capture responses
            for para in doc.paragraphs:
                text = para.text.strip()
                if text.endswith('?'):  # If the paragraph ends with a question mark, it's a question
                    current_question = text
                elif current_question and text:  # If it's not a question, it must be part of the response
                    responses.append({'Response': text, 'Name': name_tag})
            
            all_responses.extend(responses)
    
    # Convert the combined responses into a DataFrame
    df = pd.DataFrame(all_responses)
    
    return df

# Example usage with a folder path
folder_path = '/path/to/your/folder'  # Replace with actual folder path containing .docx files

# Extract feedback and save the DataFrame
combined_feedback_df = extract_feedback_from_folder(folder_path)

# Save the DataFrame to a CSV file
combined_feedback_df.to_csv("combined_feedback.csv", index=False)


import os
import pandas as pd
from docx import Document

def write_dataframe_to_word(doc_path, dataframe):
    """
    Overwrites the content of a Word document with a table based on the given DataFrame.
    """
    # Create a new Document object (overwriting the existing one)
    doc = Document()

    # Add a table with as many rows as there are dataframe rows + 1 (for headers) and columns
    table = doc.add_table(rows=dataframe.shape[0] + 1, cols=dataframe.shape[1])

    # Add headers to the first row of the table
    for i, column_name in enumerate(dataframe.columns):
        table.cell(0, i).text = column_name

    # Populate the table with DataFrame values
    for row_index, row in dataframe.iterrows():
        for col_index, value in enumerate(row):
            table.cell(row_index + 1, col_index).text = str(value)

    # Save the document, overwriting the previous content
    doc.save(doc_path)

# Example usage with a sample DataFrame
data = {
    "Name": ["Alice", "Bob", "Charlie"],
    "Feedback": ["Great!", "Needs work.", "Excellent."]
}
df = pd.DataFrame(data)

# Path to the output Word file
output_doc_path = "output_feedback_table.docx"

# Overwrite the Word document with the DataFrame's content as a table
write_dataframe_to_word(output_doc_path, df)
